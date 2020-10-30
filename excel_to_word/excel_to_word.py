from collections import namedtuple
import pandas as pd
import numpy as np
import os
import os.path
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
import progressbar
import re
import config as CONFIG


def create_files_list():
    """
    Функция создает 2 списка и Название файла (для сохранения результата). Первый список содержит пути к Эксель файлам, Второй список содержит
    пути к соответствующему данному эксель файлу (из первого списка) Ворд Шаблон
    Для сопоставления ворд и эксель файлов у них должны быть одинаковые имена
    Для ворд файлов возможены 4 варианты путей, потому что могут быть использованы  шаблоны с 5/4 колонками/таблицами и
    с сносками/без сносок
    :return: Возвращает Список списков
    """
    file_list_result = []       #Результирующий список
    current_folder = os.getcwd()
    exec_folder = os.path.join(current_folder, CONFIG.EXCEL_FOLDER) #Пути содаржания Эксель файлов
    if CONFIG.USE_A_FOOTER:
        docx_folder = os.path.join(current_folder, CONFIG.DOCX_FOLDER_4_COL) if CONFIG.FOUR_COLUMNS  else os.path.join(current_folder, CONFIG.DOCX_FOLDER_NORMAL)  # Пути содаржания ворд файлов c колонтитулом и без
    else:
        docx_folder = os.path.join(current_folder, CONFIG.DOCX_FOLDER_NO_FOOTER_4_COL) if CONFIG.FOUR_COLUMNS else os.path.join(current_folder, CONFIG.DOCX_FOLDER_NO_FOOTER)  # Пути содаржания ворд файлов c колонтитулом и без

    xlsx_files = sorted(os.listdir(exec_folder))  # Список файлов в Эксель папке
    doc_files = sorted(os.listdir(docx_folder))  # Списов файлов в ворд папке

    for x_file in xlsx_files:  # Перибираем Эксель файлы
        x_file_name, x_extension = os.path.splitext(x_file)
        for d_file in doc_files:  # Перибираем Ворд файлы
            d_file_name, d_extension = os.path.splitext(d_file)
            if x_file_name == d_file_name:  # Если названия файлов совпадат (только названия без расширений), то полные пути к этим файлам и общее название добавляютися в итоговый список
                file_list_result.append(
                    [os.path.join(exec_folder, x_file), os.path.join(docx_folder, d_file), x_file_name])
    return file_list_result


def give_table_num_list(data):
    """
    Функция, формирует список с номерами строк в ворд шаблоне, где требуется изменить номер таблицы (num_paragraphs) и
    создает новые номера для этих таблиц для замены.
    Логика: Если в конфиге указано для данного типа файлов замена номеров таблиц, то формируется два новых списка с
    номерами строк в ворде где указаны номера таблиц, и    список новых номеров для замены.
    Если в конфиге указано, что для данного вида файлов менять номера таблиц не требуется, то
    формируется два списка, равных по длине количетву таблиц в ворд файле, но заполненными значением None
    :param data: кортеж со всеми данными ( характеристика кортежа описана в функции dividing_data_into_parts)
    :return: списко номеров строк в ворде где указаны номера таблиц, список новых номеров для замены
    data.filename_for_save  - имя фала (ворд/эксель), которые сейчас обрабатываются
    data.word_data.paragraphs - все строки из вордовского файла
    CONFIG.CITY_TABLE_START_NUM  стартовый номер таблиц. От этого номера формируется список номеров таблиц в
    соответствии с количеством таблиц, например, стартовый номер 1.5, талиц 5, список равен [1.5,1.6,1.7,1.8,1.9]
    """
    start_num = False #Иницилизация переменной
    if 'city' in data.filename_for_save and 'source' not in data.filename_for_save and CONFIG.CITY_TABLE_NUM_EDIT:
        start_num = CONFIG.CITY_TABLE_START_NUM #определяем стартовый номер для номеров таблиц
    elif 'city' in data.filename_for_save and 'source' in data.filename_for_save and CONFIG.CITYSOURCE_TABLE_NUM_EDIT:
        start_num = CONFIG.CITYSOURCE_TABLE_START_NUM #определяем стартовый номер для номеров таблиц
    elif 'sklad' in data.filename_for_save and 'source' not in data.filename_for_save and CONFIG.SKLAD_TABLE_NUM_EDIT:
        start_num = CONFIG.SKLAD_TABLE_START_NUM    #определяем стартовый номер для номеров таблиц
    elif 'sklad' in data.filename_for_save and 'source' in data.filename_for_save and CONFIG.SKLADSOURDE_TABLE_NUM_EDIT:
        start_num = CONFIG.SKLADSOURDE_TABLE_START_NUM

    if start_num:   #Если стартовый номер не определен, то делаем списки заполненные None, Иначе находим строки с номерами таблиц и формируем список номеров
        num_paragraphs = [num for num, i in enumerate(data.word_data.paragraphs) if i.text.find('Таблица ') != -1]  # Номера строк с номером таблицы
        split_start_num = [int(i) for i in str(start_num).split('.')]
        paragraphs_values = [f'{split_start_num[0]}.{split_start_num[1]+i}' for i in range(len(data.word_all_tables))]  # список новых номеров для замены
        # paragraphs_values = ['{:.1f}'.format(start_num + (i / 10)) for i in range(len(data.word_all_tables))]       # список новых номеров для замены
    else:
        num_paragraphs = paragraphs_values =  [None for i in range(len(data.word_all_tables))]
    return num_paragraphs, paragraphs_values


def dividing_data_into_parts(excel_file_path, word_file_path, file_name_for_save):
    """
    Функция генератор.
    Функция создает именованный кортеж, в который записываются все необходимые для работы данные. Перечень
    характеристик  пояснены ниже при инициализации кортежа data. Кортеж создан, чтобы далее в функциях пробрасывать
    только экземпляр данного кортежа и не тянуть портянку характеристик. Кортеж создается для каждой таблицы в файле.
    Так как в одном ворд файле может быть более одной таблицы, то кортеж создается столько раз, сколько таблиц в файле,
    но функция возвращает за раз только 1 экземпляр кортежа работая как генератор (через yield по обращению цикла из
    функции manage). Генератор сделан с той целью, чтобы не формировать здесь список кортежей для всех таблиц
    т.к. необходимо будет инициализировать каждый кортежей,по числу таблиц, например 5. Если не инициализировать эти 5
    кортежей, а только один  записовать его в итоговый список, то в списке  экземпляры кортежей затрут друг друга
    последним вариантом кортежа для последней таблицы.
    Характеристика word_data не перезаписывается для каждого кортежа (т.к. генератор продолжает свою работу с места
    где закончил, т.е. в конечном цикле данной функции, в итоге  word_data редактируется скриптом столько раз, сколько
    таблиц в файле и записывается в итоговый файл.
    :param excel_file_path: Путь к Excel файлу
    :param word_file_path:  Путь к word файлу
    :param file_name_for_save: имя файла для сохранения
    :return - yield:    именованный кортеж с данными для создания и редактирования таблиц в ворд файле.
    """
    data = namedtuple('data',               #иницилизация
                      ['excel_full_file',   #содержимое эксель файла (DateFrame)
                       'excel_sheet'        #название вклдаки в эксель файле для текущей таблицы
                       'word_data',         #Полностью загруженный  Ворд файл
                       'word_all_tables'    #перечень таблиц ворд файла с контентом в формате модуля Docx
                       'select_doc_table'   #номер текущей ворд таблицы  в шаблоне для редактирования
                       'doc_table_num',     #номер строки в word_data где указан номер таблицы (Например "Таблица 1.2")
                       'doc_table_value',   #новый номер таблицы для замены ( в случаи если  номер таблицы подлежит замене)
                       'num_paragraphs_for_edit_date',  #номер строки в word_data где необходимо заменить дату
                       'num_paragraphs_for_edit_item',  #номер строки в word_data где необходимо заменить пункт (например "пункт 4.3.2")
                       'filename_for_save',             #имя файла для сохранения
                       'source'])                       # параметр False/True указывающий является ли данный Word/excel шаблон о источниках (source)

    data.word_data = Document(word_file_path)                # Загрузка текущего Ворд файла
    data.excel_full_file = pd.ExcelFile(excel_file_path)     # Загрузка текущего Эксель файла
    excel_sheets_list = data.excel_full_file.sheet_names     # Список вкладок в эксель файле
    data.word_all_tables = data.word_data.tables             # Список всех таблиц в ворд файле с содержимым
    data.filename_for_save = file_name_for_save              # имя файла для сохранения
    data.source = True if 'source' in excel_file_path else False    # параметр является ли данный файл о источчниках парсинга
    doc_table_num, doc_table_value = give_table_num_list(data)      #Получить списки номеров строк где указан номер таблицы и новые номера
    data.row_not_for_replace = 3 if data.source else 2  # Количество строк в таблице не подлежащих редактированию (Заголовки)
    if data.source:     # Если файл о источниках то формируются списки номеров строк где указана дата, и списки номеров строк где указан пункт
        num_paragraphs_for_edit_date = [num for num, i in enumerate(data.word_data.paragraphs) if i.text.find('предоставленных в дату предоставления') != -1]  # Номера строк с датами
        num_paragraphs_for_edit_item = [num for num, i in enumerate(data.word_data.paragraphs) if i.text.find('4.3.') != -1]  # Номера строк с пунктами
    else: #Если файл не о источниках то  формируются списки с значением None, вданных файлах не редатируются даты и пункты с троках (не таблицах)
        num_paragraphs_for_edit_date = num_paragraphs_for_edit_item = [None]

    # #Этот блок нужен для отладки, по хорошему нужно удалить, но код дорадатывается, постоянно это писать не хочу.
    # print('данные ворд ', data.word_data)
    # print('данные эксель ', data.excel_full_file)
    # print('Вкладки эсель ',excel_sheets_list)
    # print('номер таблицы для замены ',doc_table_num)
    # print('номер таблицы на который заменить ',doc_table_value)
    # print('номер параграфа для замены даты ',num_paragraphs_for_edit_date)
    # print('номер параграфа для замены пункта ',num_paragraphs_for_edit_item)
    # print('имя файла ля соханения',data.filename_for_save)
    # Проверка что колличество характеристик строк в которые нужно вносить изменения равно количеству таблиц
    assert len(data.word_all_tables) == len(doc_table_num) == len(doc_table_value) == len(num_paragraphs_for_edit_date) == \
           len(num_paragraphs_for_edit_item) == len(excel_sheets_list), 'Количество таблиц в ворде, вкладок (таблиц)' \
            ' в экселе,  параграфов замены даты, парафов замены пунктов, номеров таблиц, и номеров для замены не совпадает'
    #Запись в кортеж персональных для каждой таблицы характеристик
    for num in range(len(data.word_all_tables)):
        data.select_doc_table = data.word_all_tables[num]   #текущая таблица
        data.excel_sheet = excel_sheets_list[num]           #текущая вкладка
        data.doc_table_num = doc_table_num[num]             #номер строки с номером таблицы
        data.doc_table_value = doc_table_value[num]         #новое значение для номера таблицы
        data.num_paragraphs_for_edit_date = num_paragraphs_for_edit_date[num]   # номер стироки для замены даты
        data.num_paragraphs_for_edit_item = num_paragraphs_for_edit_item[num]   # номер строки для замены пункта
        yield data  #вернем экземпляр кортежа


def load_and_prepare_excel(excel_info, sheet, source):
    """
    Функция загружает таблицу из эксель файла excel_info расположенной на вкладке sheet
    Удаляет из эксель таблицы не нужные столбцы, форматирует эксель таблицу так, чтобы  ее можно было вставить в Ворд таблицу, это
    Для файлов sourc:
        Последняя строка состоит из 3-х ячеек, первая из которых это 9-ть объединенных ячеек ,
        чтобы последняя строка выглядела как |Итого| 0000| 0000| в датафрейме нужно сделать в последней строке
        вот так |не важно|не важно|не важно|не важно|не важно|не важно|не важно|не важно|Итого:|0000|0000|
    Для других файлов добавляет строку с суммой по каждому столбцу с числовыми данными (в эксель таблице этих данных нет а в Ворд таблице они нужны)
    :param excel_info: Общий файл эксель
    :param sheet:   Название вкладки в эксель файле
    :param source: True/False Если мы работаем с файлом/данными "источников", например файл city_source.docx
    :return:  Возвращает эксалевскую таблицу в виде  pandas DateFrame и параметр equal_result (True/False),
     который показывает, нужно ли менять пункт указанные в тексте перед таблицей по условию.
    """
    excel_data = pd.read_excel(excel_info, sheet_name=sheet).fillna(0) # Загрузка таблицы с конкретной вкладки файла
    col_name_for_del = ['subject', 'city', 'Unnamed: 0']  # Столбцы, которые требуется удалить если они есть в таблице
    equal_result = False
    for col_name in col_name_for_del:  # Удаление ненужных столбцов если они есть эксель данных
        if col_name in excel_data.columns:
            excel_data = excel_data.drop([col_name], axis=1)
    if source:  # в описании данной функции поеснино (для файлов sourc)
        excel_data.iloc[-1, 9] = 'Итого:'
        equal_result = True if excel_data.iloc[-1][-1] == excel_data.iloc[-1][-2] else False  # если данные в двух последних ячейках последней строки совпадают то True иначе False
    else:  # в описании данной функции поеснино (для файлов sourc)
        for col_for_type in excel_data.columns:
            excel_data[col_for_type] = excel_data[col_for_type].astype('int')

        #excel_data.number = excel_data.number.astype('int')  # Изменить тип столбца на целочисленный, потому что в эксель он указан строкой,  и поэтому неверно выравнивается в ячейке Ворд таблицы.
        total_row = excel_data.sum()                        # Строка с суммами по всем строкам таблицы
        total_row['number'] = 'Итого:'                      # Меняем первую ячсейку в соответстви с требованиями Ворд шаблона
        excel_data = excel_data.append(total_row, ignore_index=True)    #В итоговую таблицу добавляем стороку с суммами по каждому столбцу

    return excel_data, equal_result


def get_cell_format(cell):
    """
    Функция сохранения параметров форматированя текста и шрифта текущей ЯЧЕЙКИ
    :param cell:  Текущая ячейка
    :return:    Возвращает параметры выравнивания, размера шрифта, названия шрифта, жирный текст или нет
    """
    run = cell.paragraphs[0].runs       #Получение настроек ячейки
    alignment = 0 if cell.paragraphs[0].alignment == None else int(
        re.findall('\d', str(cell.paragraphs[0].alignment))[0])
    font = run[0].font      #Получение характеристик шрифта ячейки
    font_size = font.size
    font_name = font.name
    font_bold = font.bold
    return alignment, font_size, font_name, font_bold


def set_cell_format(cell, font_ali, font_size, font_bold, font_name):
    """
    Фунция применяет к текущей ЯЧЕЙКИ параметры форматирования выравнивания, размера и тд.
    :param cell:    текущая ячейка
    :param font_ali:    выравнивание    по гооризонтали
    :param font_size:  размер текста
    :param font_bold:   Жирный/нежирный
    :param font_name:   Название шрифта
    :return:    ничего не возвращает так как изменяет параметры ячейки по ссылке cell
    """
    run = cell.paragraphs[0].runs #Получение настроек ячейки
    cell.paragraphs[0].alignment = font_ali  # 0 for left, 1 for center, 2 right, 3 justify ....
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Выравнивание текста по центру по ВЕРТИКАЛИ
    font = run[0].font
    font.size = font_size
    font.bold = font_bold
    font.name = font_name


def get_parag_format(parag):
    """
    Функция сохранения параметров форматированя текста и шрифта текущей СТРОКИ
    :param parag: текущая строка
    :return: Возвращает параметры выравнивания, размера шрифта, названия шрифта, жирный текст или нет
    """
    run = parag.runs    #Получение настроек  строки
    font = run[0].font  #Получение характеристик шрифта строки
    font_size = font.size
    font_name = font.name
    font_bold = font.bold
    return font_size, font_name, font_bold


def set_parag_format(parag, font_size_was, font_name_was, font_bold_was):
    """
     Фунция применяет к текущей СТРОКЕ параметры форматирования текста
    :param parag:  текущая строка
    :param font_size_was:  размер текста
    :param font_name_was:   Название шрифта
    :param font_bold_was:  Жирный/нежирный
    :return: ничего не возвращает так как изменяет параметры текста по ссылке parag
    """
    run = parag.runs    #Получение настроек  строки
    font = run[0].font  #Получение характеристик шрифта строки
    font.size = font_size_was
    font.bold = font_bold_was
    font.name = font_name_was


def edit_date_and_item_in_parag(current_table,target=''):
    """
    Функция редактирует дату и пункт (если требуется) на который ссылается текст перед текущей таблицей
    :param current_table: Кортеж со всеми характеристиками для изменения табл
    :param target:  Цель, что будем менять, дату в  строке, пункт в строке, или номер таблицы в строке
    :return: ничего не возвращает так как редактирует  строку (текст перед таблицей), по ссылке  в кортеже
    parag = current_table.word_data.paragraphs[current_table.num_paragraphs_for_edit_date]
    Пояснение Из общего кортежа с данных (current_table)  из списка строк текущей таблицы (current_table.word_data.paragraphs)
    по номеру выбирается строка (current_table.num_paragraphs_for_edit_date) содержащая дату.
    """
    if target == 'date':  #Замена даты в строке
        parag = current_table.word_data.paragraphs[current_table.num_paragraphs_for_edit_date] # выбор строки по номеру из списка строк
        format_rule = get_parag_format(parag)   #Сохранение параметров форматированя текста и шрифта
        date_in_parag = re.findall(r'\d{2}\.\d{2}\.\d{4}', parag.text)[0]                   #Находим даты в тексте
        parag.text = parag.text.replace(date_in_parag, current_table.excel_sheet)           #Заменяем дату на новую
        set_parag_format(parag,*format_rule)    #Применить параметры форматированя текста и шрифта
    elif target == 'item':  #Замена пункта в строке
        parag = current_table.word_data.paragraphs[current_table.num_paragraphs_for_edit_item]
        format_rule = get_parag_format(parag)    #Сохранение параметров форматированя текста и шрифта
        parag.text = parag.text.replace('4.3.2', '4.3.3')                  # Заменяем дату на новую
        set_parag_format(parag, *format_rule)   #Применить параметры форматированя текста и шрифта
    elif target == 'table_num': #Замена номера таблицы в строке
        parag = current_table.word_data.paragraphs[current_table.doc_table_num]
        format_rule = get_parag_format(parag)    #Сохранение параметров форматированя текста и шрифта
        replace_it = re.findall(r'[\d.]{3,5}', parag.text)[0]
        parag.text = parag.text.replace(replace_it, current_table.doc_table_value)  # Заменяем номер на новый
        set_parag_format(parag, *format_rule)   #Применить параметры форматированя текста и шрифта


def edit_head_in_doc_table(doc_table, excel_table):
    """
    Функция редактирует даты в заголовке Ворд таблицы на те, что присутствуют в заголовках Экселесвкой таблицы.
    Функция редактирует пункты ('4.3.2', '4.3.3') в заголовке. Условие, в одну дату, если совпадают суммы по столбцам
    "Фактическое количество" и Статус актуальной ссылки", тогда пункт '4.3.3', иначе '4.3.2'
    В данной функции используется list.pop() для перемещения по спискам, чтобы постепенно выбирать элемент за элементом
    а не делать  какие-то указатели на соответствующий индекс элемента.
    :param doc_table:       Вордовская таблица
    :param excel_table:     Экселевская таблица
    :return:  Ничего не возвращает т.к. редактирует таблицу  по ссылке doc_table
    """
    pat_date = r'[\d.]{10}'  # патерн поиска даты для регулярки
    # из списка колонок выбирает колонки где есть даты (регуляркой)
    date_from_excel_columns = [re.findall(pat_date, i)[0] for i in excel_table.columns if re.findall(pat_date, i)][::-1] #развернуть для pop()
    condition_list_edit_par =  []           #Список False/True для замены пункта в заголовке ил инет.
    condition_data = excel_table.iloc[-1][1:]       #берем последниюю строку таблицы с суммами по столбцам, исключая первую колонку, там написано "Итого"
    for item in np.array_split(condition_data, len(condition_data) / 2):    #Получившийся массив делим по два столбца (Фактические и актуальные) и перечисляем его
        for _ in range(2): condition_list_edit_par.append(True if np.unique(item).size == 1 else False) #Если значения двух элементов списка совпадают (т.е. совпадают фактические и актуальные)
        # то дважды добавляем True (Совпадают) или False (Несовпадают). Дважды, потому что в заголовке 10/8 столбцов с датами и пунктами, содениненых ПО ДВА, но фактически их 10/8
    condition_list_edit_par = list(reversed(condition_list_edit_par)) # Разварачиваю список чтобы выбирать значения спомощью list.pop()

    doc_cols = len(doc_table.columns)  # Количество колонок в Вордовской таблице
    num_row_for_edit = 0  # первая строка, которую необходимо править, там где даты
    for col in range(doc_cols):  # перебираю все колонки
        cell = doc_table.rows[num_row_for_edit].cells[col]  # выбор конкретной ячейки
        date_in_cell = re.findall(pat_date, cell.text)  # Поиск в тексте ячейки даты
        if date_in_cell:  # Если дата есть
            font_ali_was, font_size_was, font_name_was, font_bold_was = get_cell_format(cell)  # Получить характеристики текста
            try:
                cell.text = cell.text.replace(date_in_cell[0],date_from_excel_columns.pop())  # Заменить дату ячейки на ту, что была в списке заголовков экселевской таблицы
            except IndexError:
                print('Число столбцов в эксель файлах, которые нужно вставить ворд не соответсвует числу столбков в шаблоне ворд,\n'
                      'или число таблиц в экселе не соответствует числу таблиц в шаблоне ворд. За эти настройки отвечают \n'
                      'параметры USE_A_FOOTER и FOUR_COLUMNS в файле конфига config.py, и необходимо проверить эксель файл с данными')
            if condition_list_edit_par.pop():   #нужно ли менять пункт в данном ячейке
                cell.text = cell.text.replace('4.3.2', '4.3.3')

            set_cell_format(cell, font_ali_was, font_size_was, font_bold_was, font_name_was)  # Применить характеристики текста

    doc_table.rows[num_row_for_edit].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY  # Разрешить редактирование высоты ячейки
    doc_table.rows[num_row_for_edit].height = Pt(CONFIG.HEIGHT_OF_CELL_IN_HEADER)  # Установить высоту ячейки


def add_row_in_doc_table(current_table,excel_row):
    """
    Изначально в шаблоне таблицы Ворд 3 или 4 строки, это 2 строки заголовка (в том числе объединенные ячейки),
    1 строка номера колонки (в файлах источников), 1 строка с данными для образца форматирования ячеек и текста.
    Функция добавляет в таблицу ворд столько строк (без строк заголовков), чтобы количество соответствовало количеству строк из эксель файла
    из которого берутся  данные для таблицы.
    :param current_table: Кортеж со всеми характеристиками для изменения таблиц
    :param excel_row: Количество строк в эксель файле
    :return: Ничего не возвращает, т.к. функция вносит изменения в таблицу по ссылкам в кортеже
    current_table.row_not_for_replace - Количество строк  уже присутствующих в таблице (строки заголовков и шаблона форматрования)
    current_table.source True/False Если мы работаем с файлом/данными "источников", например файл city_source.docx, в таком случаи
    последнуу строку итогов нужно объединить, чтобы получилось 3 ячейки.
    """
    doc_rows = len(current_table.select_doc_table.rows)                      #текущее кол-во строк в Вод таблице
    doc_cols = len(current_table.select_doc_table.columns)                   #текущее количество столбцов в таблице
    delta = excel_row - doc_rows + current_table.row_not_for_replace         #Колчество строк которое необходимо добавить, чтобы в итоговой таблице было столькоже строк с данными сколько в экселе
    for _ in range(delta):
        current_table.select_doc_table.add_row()                             #Добавляем строки
    if current_table.source:                                                 #Если тип файла source, в последней строке  объединяем ячейки, чтобы получилось 3 колонки
        for col in range(doc_cols): #перебрать кольнки
            if col < doc_cols - 2:  #Все колонки кроме последних двух объдиняяяяяяяем
                current_table.select_doc_table.rows[-1].cells[0].merge(current_table.select_doc_table.rows[-1].cells[col])


def save_result_docx(doc_data, filename):
    """
    Функция Сохраняет итоги соединения эксель файла и вордовского шаблона в папку RESULT_DOCX.
    Если такая папка не существовала рядом с нашим скриптом, то создаем ее.
    :param doc_data: Итоговые данные для сохранения
    :param filename: имя файла для сохранения
    :return: Нечего возвращать
    """
    current_folder = os.getcwd()        #Текущая папка
    result_folder = os.path.join(current_folder, 'RESULT_DOCX') #Куда хотим сохранить
    if not os.path.exists('RESULT_DOCX'):   #Если такой папки нет, создаем
        os.mkdir(result_folder)
    doc_data.save(os.path.join(result_folder, f'{filename}.docx'))  #Сохраняем
    print(f'* файл {filename} сохранен')


def join_exceldata_and_docxtable(current_tab):
    """
    1) Загружает и подготавливает  текущие ворд и эксель таблицы для соединения.
    Эксель данные  подготавливаются в функции load_and_prepare_excel
    В зависимости от типа файла (source?) и настроек в конфиг файле редактируются строки перед таблицами с датами,
    пунктами, номерами таблиц в функции edit_date_and_item_in_parag.
    Ворд таблица подготавливаются с помощью добавления строк в функции add_row_in_doc_table
    2)Соединяет таблицы с учетом форматирования ячеек и строк
    3)Отображает прогресс бар в терминале по ходу процесса заполнения шаблонов таблиц Ворд данными из экселя
    :param current_tab: Кортеж со всеми характеристиками для изменения таблиц для конкретной таблицы,
                        все элементы описаны в функции dividing_data_into_parts
    :return: ничегоне возвращает, просто меняет данные в кортеже (который доступен по всей программе исохранится в файл в другой функции)
    edit_paragraph_item_bool  - True/false менять ли строку перед таблицей с пунктом, определяется в
    функции load_and_prepare_excel
    current_tab -  doc_table_num и doc_table_value заполнены если требуется внести изменения иначе None
    """

    excel_data, edit_paragraph_item_bool = load_and_prepare_excel(current_tab.excel_full_file, current_tab.excel_sheet, current_tab.source) #Загруженная и подготовленая конкретная эксель таблица, и значение необходимости менять пункт параграфа
    excel_row, excel_col = excel_data.shape #количество строк и столбцов в эксель, чтобы перебирать, как всегда.
    items_row = current_tab.select_doc_table.rows[-1]  # Последняя строка шаблона ВОРД  как образец параметров выравнивания, имени шрифта , размера шрифта при формировании таблицы
    if current_tab.num_paragraphs_for_edit_item and edit_paragraph_item_bool: #в описании данной функции
         edit_date_and_item_in_parag(current_tab,target ='item')
    if current_tab.num_paragraphs_for_edit_date:
        edit_date_and_item_in_parag(current_tab,target ='date') #Редактирование даты в строке перед таблицей
    if current_tab.doc_table_value: #в описании данной функции
        edit_date_and_item_in_parag(current_tab, target='table_num')  # Редактирование даты в строке перед таблицей
    if not current_tab.source:
        edit_head_in_doc_table(current_tab.select_doc_table,excel_data)

    add_row_in_doc_table(current_tab,excel_row) #добавление в шаблон Ворда нужное количество строк (подробнее в функции add_row_in_doc_table)
    rows = len(current_tab.select_doc_table.rows)  #Количество строк в ворд таблице после  добавления строк
    cols = len(current_tab.select_doc_table.columns)   #Количество колонок в  ворд таблице
    bar = progressbar.ProgressBar(max_value=rows)   #Для отображения динамики происходящего

    for row in range(rows): #Перебираем строки ворд таблицы
        bar.update(row)     #Для отображения динамики происходящего
        if row - current_tab.row_not_for_replace < excel_row:   #В ворд таблице строк больше чем в Эксель, так как в ворде заголовки считаются строками (0 в итоге в ворде на 2-3 строки больше в зависимости от файла
            for col in range(cols):                 #Перебираем колонки
                cell = current_tab.select_doc_table.rows[row].cells[col]   #Текущая ячейка
                items_cell = items_row.cells[col]               #Соответствующая ячейка из строки шаблона, специально оставленная для  определения необходимого форматирования всех новых строк в шаблоне
                if row > current_tab.row_not_for_replace - 1:  # Если строки не заголовок (ниже чем row_not_for_replace)
                    font_ali_was, font_size_was, font_name_was, font_bold_was = get_cell_format(items_cell)  # Получить характеристики текста
                    cell.text = str(excel_data.iloc[row - current_tab.row_not_for_replace][col]) #Заменить текст в ячеки соотвествующим текстом Эксель ячейки
                    if row == rows - 1: #Если строка последняя и файл типа source, то у него Должен быть жирный шриф
                        font_bold_was = True if current_tab.source else font_bold_was
                        font_ali_was = 0 if col < cols - 2 and current_tab.source else font_ali_was #А если .то первая ячейка последней строки, то у нее выравнивание по левому краю
                    set_cell_format(cell, font_ali_was, font_size_was, font_bold_was,font_name_was)  # Применить характеристики текста

                    current_tab.select_doc_table.rows[row].height_rule = None if current_tab.source else WD_ROW_HEIGHT_RULE.EXACTLY    # Разрешить редактирование высоты ячейки Если файл source
                    current_tab.select_doc_table.rows[row].height = Pt(CONFIG.HEIGHT_OF_CELL_IN_SOURCE) if current_tab.source else Pt(CONFIG.HEIGHT_OF_CELL_IN_NOT_SOURCE) # Установить высоту ячейки, в зависимости от параметра source
    bar.finish()    #Для отображения динамики происходящего


def manage():
    """
    Функция управления.
    1) Создает список  содержащий пути к Эксель файлам, соответствующие им  Ворд файлы и названяи итоговых файлов
    для сохранения (подробнее в create_files_list)
    2) перебирает данный список передавая пути файлов в генератор dividing_data_into_parts
    3) dividing_data_into_parts  создает именованный кортеж содержащий данные о таблицых (подробнее в dividing_data_into_parts)
    4) Поскольку таблиц в каждом файле может быть более 1 то из генератора dividing_data_into_parts через цыкл
    перебираются  кортежи с данными для редактирования каждой таблицы. Непосредственно редактирование каждой таблицы и
    подготовка данных осуществляется в  функции join_exceldata_and_docxtable.
    содержимое итогового ворд файла сохраняется через функцию save_result_docx
    :return: Ничего не возвращает
    """
    files_path_list = create_files_list() #Создать список путей к файлам шаблонов и данных
    for excel_file_path, word_file_path, file_name_for_save in files_path_list: #Перебираем этот список постепенно обрабатывая все файлы
        print('* {} + {} ---> {}.docx'.format(os.path.split(excel_file_path)[1], os.path.split(word_file_path)[1], file_name_for_save))
        table_list = dividing_data_into_parts(excel_file_path,word_file_path,file_name_for_save)
        for table_data in table_list:
            join_exceldata_and_docxtable(table_data)
            word_data = table_data.word_data
        save_result_docx(word_data, file_name_for_save)  # Сохраняем результаты редактиирования Ворд фала



if __name__ == '__main__':
    manage()